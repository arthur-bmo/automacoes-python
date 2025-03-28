import pandas as pd
import docx
from docx2pdf import convert
from docx import Document
import locale

def substituir_texto(documento, marcador, texto_substituto):
    for p in documento.paragraphs:
        if marcador in p.text:
            p.text = p.text.replace(marcador, texto_substituto)
    for tabela in documento.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                if marcador in celula.text:
                    celula.text = celula.text.replace(marcador, texto_substituto)

locale.setlocale(locale.LC_ALL, 'pt_BR.UTF-8')

def format_brazilian_real(value):
    try:
        # Convert the string to a float
        number = float(value)
        # Format the number as currency
        return locale.currency(number, grouping=True)
    except ValueError:
        return "Invalid number"

def gerar_documento(dados):
    modelo = Document('A:/Trabalho/EMBRAPII/Programa Lívia/Modelo recibo.docx')
    for _, linha in dados.iterrows():
        try:
            new_path = str(f'{linha["Bolsista"]} {linha["Mês"]} {linha["Ano"]}.docx')
            modelo.save(new_path)
            doc_novo = Document(new_path)  # Clonar o documento modelo
            substituir_texto(doc_novo, 'NBOLSISTO', linha['Bolsista'])
            substituir_texto(doc_novo, 'CEPEEFE', str(linha['CPF']))
            substituir_texto(doc_novo, 'ERREGE', linha['RG'])
            substituir_texto(doc_novo, 'INDERECU', linha['Endereço'])
            valor_format = format_brazilian_real(linha['Valor'])
            valor_format = str(valor_format)
            substituir_texto(doc_novo, 'VALUEDABOLSA', valor_format)
            substituir_texto(doc_novo, 'DATAQCONFIRMA', str(linha['Data de confirmação']))
            substituir_texto(doc_novo, 'AREFERENCIA', str(linha['Referência']))
            substituir_texto(doc_novo, 'OMES', linha['Mês'])
            substituir_texto(doc_novo, 'OANO', str(linha['Ano']))
            substituir_texto(doc_novo, 'PROJETODAFUNARBE', linha['CodigoProjeto'])
            substituir_texto(doc_novo, 'NOMEDOPROJETO', linha['ProjetoFunarbe'])
            hora_format = linha['Horas mensais']
            hora_format = str(hora_format).replace(".0","")
            substituir_texto(doc_novo, 'COORDOPROJETO', linha['Coordenador'])
            
            substituir_texto(doc_novo, 'QDEHORAS', hora_format)
            
            doc_novo.save(new_path)
            print(f"Arquivo salvo: {new_path}")
            novo_pdf = convert(new_path)

        except Exception as e:
            print(f'Erro: {e}')

if __name__ == "__main__":
    # Carregar os dados da planilha
    dados = pd.read_excel('planilha info recibos.xlsm', sheet_name = 'dados corretos')
    
    # Gerar documentos com base no modelo e dados
    documento = Document('Modelo recibo.docx')
    gerar_documento(dados)
